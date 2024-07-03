from docx import Document
from googletrans import Translator

def translate_docx(input_path, output_path, src_lang='en', dest_lang='kn'):
    # Load the document
    document = Document(input_path)
    translator = Translator()

    # Translate text in each paragraph
    for paragraph in document.paragraphs:
        if paragraph.text.strip():  # Only translate non-empty paragraphs
            print("original:", paragraph.text)
            translated_text = translator.translate(paragraph.text, src=src_lang, dest=dest_lang).text
            print("translated:", translated_text)
            paragraph.text = translated_text

    # Translate text in each table cell
    #for table in document.tables:
    #    for row in table.rows:
    #        for cell in row.cells:
    #            if cell.text.strip():  # Only translate non-empty cells
    #                print("original:", cell.text)
    #                translated_text = translator.translate(cell.text, src=src_lang, dest=dest_lang).text
    #                print("translated:", translated_text)
    #                cell.text = translated_text

    # Save the translated document
    document.save(output_path)

# Example usage
input_docx = 'documents/demo1.docx'
output_docx = 'documents/output1demo1.docx'
translate_docx(input_docx, output_docx, src_lang='en',dest_lang='kn')