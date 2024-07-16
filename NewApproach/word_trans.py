from docx import Document
from googletrans import Translator

def translate_paragraph(paragraph, translator, src_lang, dest_lang):
    # Translate the text of the paragraph
    if paragraph.text.strip():
        translated_text = translator.translate(paragraph.text, src=src_lang, dest=dest_lang).text
        return translated_text
    return paragraph.text

def translate_docx(input_path, output_path, src_lang='en', dest_lang='kn'):
    # Load the document
    document = Document(input_path)
    translator = Translator()

    # Create a new document to preserve the structure
    new_document = Document()

    # Translate text in each paragraph and keep the formatting
    for paragraph in document.paragraphs:
        new_paragraph = new_document.add_paragraph()
        for run in paragraph.runs:
            if run.text.strip():  # Check if run text is not empty
                translated_text = translator.translate(run.text, src=src_lang, dest=dest_lang).text
                new_run = new_paragraph.add_run(translated_text)
                # Preserve the formatting
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size

    # Translate text in each table cell
    for table in document.tables:
        new_table = new_document.add_table(rows=0, cols=len(table.columns))
        for row in table.rows:
            new_row = new_table.add_row()
            for cell_index, cell in enumerate(row.cells):
                new_cell = new_row.cells[cell_index]
                for paragraph in cell.paragraphs:
                    translated_text = translate_paragraph(paragraph, translator, src_lang, dest_lang)
                    new_paragraph = new_cell.add_paragraph(translated_text)

    # Save the translated document
    new_document.save(output_path)

# Example usage
input_docx = 'documents/demo1.docx'
output_docx = 'documents/output1demo1.docx'
translate_docx(input_docx, output_docx, src_lang='en', dest_lang='kn')
