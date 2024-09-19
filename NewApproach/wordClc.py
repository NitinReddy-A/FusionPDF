from docx import Document

def remove_paragraphs_by_index(file_path, output_path, para_indices_to_remove):
    # Load the document
    doc = Document(file_path)

    # Sort the indices in reverse order so removing paragraphs doesn't affect the numbering
    para_indices_to_remove = sorted(para_indices_to_remove, reverse=True)

    # Loop through the paragraphs and remove based on their indices
    for index in para_indices_to_remove:
        if 0 <= index < len(doc.paragraphs):
            # Clear the paragraph's text at the specified index
            para = doc.paragraphs[index]
            para.clear()

    # Save the modified document
    doc.save(output_path)

# Example usage:
file_path = 'output.docx'
output_path = 'FinOp.docx'
para_indices_to_remove = [0,1,2,3,4,5,6,7,23,24]  # Example: Remove paragraphs 2, 4, and 6 (0-based index)

remove_paragraphs_by_index(file_path, output_path, para_indices_to_remove)
