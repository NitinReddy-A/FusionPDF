from docx import Document
import requests
import os
import time
import fitz
import pymupdf
from PIL import Image
import json
from fpdf import FPDF
import convertapi
from googletrans import Translator

def create_ocr_task(api_key, file_path, output_format='docx'):
    """
    Creates an OCR task for a scanned document.
    
    :param api_key: The API key for authentication.
    :param file_path: The path to the scanned PDF file.
    :param output_format: The desired output format (default is 'docx').
    :return: The task ID if successful, otherwise an error message.
    """
    OCR_CONVERSION_URL = 'https://techhk.aoscdn.com/api/tasks/document/ocr'
    
    # Prepare the request headers and payload
    headers = {'X-API-KEY': api_key}
    files = {'file': open(file_path, 'rb')}
    data = {'format': output_format}
    
    try:
        # Send a POST request to create the OCR conversion task
        response = requests.post(OCR_CONVERSION_URL, headers=headers, data=data, files=files)
        
        # Check if the request was successful
        if response.status_code == 200:
            task_id = response.json().get('data', {}).get('task_id')
            if task_id:
                return task_id
            else:
                return f"Error: {response.json().get('message')}"
        else:
            return f"Error: {response.status_code} - {response.text}"
    except Exception as e:
        return f"Exception occurred: {str(e)}"


def perform_ocr_task(task_id, api_key, result_url, output_path='o.docx', timeout=30):
    """
    Perform OCR task by checking status and downloading the result.

    :param task_id: ID of the OCR task
    :param api_key: API key for authentication
    :param result_url: URL endpoint for the OCR task
    :param output_path: Path to save the output file
    :param timeout: Maximum time in seconds to wait for task completion
    """
    headers = {'X-API-KEY': api_key}
    
    for _ in range(timeout):
        # Check OCR task status
        response = requests.get(f"{result_url}{task_id}", headers=headers)
        if response.status_code != 200:
            print(f"Error checking OCR status: {response.status_code} - {response.text}")
            return False

        result = response.json()
        state = result.get('data', {}).get('state')

        if state == 1:  # State 1 indicates completion
            file_url = result.get('data', {}).get('file')
            if file_url:
                # Download the file
                file_response = requests.get(file_url)
                if file_response.status_code == 200:
                    with open(output_path, 'wb') as output_file:
                        output_file.write(file_response.content)
                    print(f"OCR conversion successful! Output saved as '{output_path}'.")
                    return True
                else:
                    print(f"Error downloading the file: {file_response.status_code} - {file_response.text}")
                    return False
            else:
                print("Error: File URL not found in the response.")
                return False
        elif state < 0:  # Any negative state indicates failure
            print(f"OCR conversion failed with state: {state}.")
            return False
        else:
            progress = result.get('data', {}).get('progress', 0)
            print(f"OCR conversion in progress... {progress}% complete.")
            time.sleep(1)  # Wait 1 second before checking again

    print("OCR conversion timed out.")
    return False


# Set your API key and file path
# API_KEY = 'wx6tkon97x0q5qyl8'
# RESULT_URL = 'https://techhk.aoscdn.com/api/tasks/document/ocr/'
# pdf_file_path = r'C:\Users\Lenovo\Desktop\repo\FusionPDF\documents\scan1.pdf'
# 
# Step 1: Create OCR Task
# task_id = create_ocr_task(API_KEY, pdf_file_path, output_format='docx')
# 
# if 'Error' not in task_id:
    # print(f"Task created successfully. Task ID: {task_id}")
    # 
    # Step 2: Retrieve OCR Result
    # result_url = perform_ocr_task(task_id=task_id, api_key=API_KEY, result_url=RESULT_URL)
    # 
    # print(f"Processed document is available at: {result_url}")
    # 
# else:
    # print(task_id)

def remove_paragraphs_by_index(file_path, output_path, para_indices_to_remove):
    # Load the document
    doc = Document(file_path)

    # Sort the indices in reverse order so removing paragraphs doesn't affect the numbering
    para_indices_to_remove = sorted(para_indices_to_remove, reverse=True)

    # Loop through the paragraphs and remove based on their indices
    for index in para_indices_to_remove:
        if 0 <= index < len(doc.paragraphs):
            # Clear the paragraph's text at the specified index
            para = doc.paragraphs[index]
            para.clear()

    # Save the modified document
    doc.save(output_path)

# Example usage:
# file_path = 'o.docx'
# output_path = 'TextOp.docx'
# para_indices_to_remove = [0,1,2,3,4,5,6,7,23,24]  # Example: Remove paragraphs 2, 4, and 6 (0-based index)
# 
# remove_paragraphs_by_index(file_path, output_path, para_indices_to_remove)
# 
# convertapi.api_credentials = 'secret_HVuqFuKW4UsSHiCI'
# convertapi.convert('pdf', {
    # 'File': r'TextOp.docx' 
# }, from_format = 'docx').save_files('ScannedPDF_Final')

def convert_pdf_to_jpg(pdf_path, output_folder, zoom=2):
    doc = fitz.open(pdf_path)
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        
        # Set a matrix to scale the page by the zoom factor (2 = 200% size)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        
        pix.save(f"{output_folder}/page_{page_num + 1}.jpg")

    doc.close()

# convert_pdf_to_jpg(r"documents\scan1.pdf", "ScannedPDF_Final",zoom=3)

def save_headers(pdf_path,left, upper, right, lower):
    # Open the image
    image = Image.open(pdf_path)

    # Define the coordinates for cropping (left, upper, right, lower)
    # left = 0
    # upper = 100
    # right = 1780
    # lower = 650

    # Crop the image
    cropped_image = image.crop((left, upper, right, lower))

    # Display or save the cropped image
    cropped_image.show()
    cropped_image.save("Header.jpg")


def save_footers(pdf_path,left1, upper1, right1, lower1):
    # Open the image
    image = Image.open(pdf_path)

    # left1 = 0
    # upper1 = 2235
    # right1 = 1780
    # lower1 = 2520

    # Crop the image
    cropped_image1 = image.crop((left1, upper1, right1, lower1))

    # Display or save the cropped image
    cropped_image1.show()
    cropped_image1.save("Footer.jpg")

# save_headersNfooters(r"ScannedPDF_Final\page_1.jpg")


def create_pdf_with_images(header_image, footer_image, output_pdf_path):
    """
    Creates a PDF document with a header and footer image.

    :param header_image: Path to the header image.
    :param footer_image: Path to the footer image.
    :param output_pdf_path: Path to save the generated PDF.
    """
    class PDFWithImages(FPDF):
        def header(self):
            self.image(header_image, x=10, y=10, w=self.w - 20)

        def footer(self):
            self.set_y(-30)
            self.image(footer_image, x=10, w=self.w - 20)

    # Initialize PDF document
    pdf = PDFWithImages()
    pdf.add_page()
    
    # Output the PDF to the specified file path
    pdf.output(output_pdf_path)

# header_image = "Header.jpg"  # Path to the header image
# footer_image = "Footer.jpg"  # Path to the footer image
# output_pdf_path = "FinalOutput.pdf"  # Path where the PDF will be saved
# 
# create_pdf_with_images(header_image, footer_image, output_pdf_path)



def extract_text_with_coordinates(pdf_path, output_json_path):
    """
    Extract text from a PDF file along with coordinates, font size, and color,
    and save the data to a JSON file.

    :param pdf_path: Path to the input PDF file.
    :param output_json_path: Path to the output JSON file.
    """
    # Dictionary to store text with coordinates and character count
    extracted_data = {}
    
    # Create a document object
    doc = fitz.open(pdf_path)
    
    # Extract the number of pages
    print(f"Number of pages: {doc.page_count}")
    
    # Extract metadata
    print("Metadata:", doc.metadata)
    
    # Iterate through all pages
    for i in range(doc.page_count):
        # Get the page
        page = doc.load_page(i)  # or page = doc[i]
        
        # Extract text blocks from the page
        blocks = page.get_text("dict", flags=11)["blocks"]
    
        #blocks1 = page.get_text("blocks")
        
        # List to store text and coordinates for the current page
        page_data = []
    
    
    
        for b in blocks:
            bbox = b["bbox"]
            for l in b["lines"]:  # iterate through the text lines
                for s in l["spans"]:  # iterate through the text spans
                    text = s["text"]
                    #origin = s["origin"]
                    font_size = s["size"]
                    #character_count = len(text)
                    color = s["color"]
    
            # Append to the result
            page_data.append({
                "coordinates": bbox,
                "text": text,
                #"IniCharacter_count": character_count,
                "IniFontsize": font_size,
                "Color": color,
            })
    
        # Add the page data to the extracted data dictionary
        extracted_data[i+1] = page_data
    
    # Save the extracted data to a JSON file
    with open(output_json_path, "w", encoding="utf-8") as json_file:
        json.dump(extracted_data, json_file, ensure_ascii=False, indent=4)
    
    # Close the document
    doc.close()
    print(f"Extraction complete. Data saved to {output_json_path}.")


def add_text_and_character_count(pdf_path, json_path):
    """
    This function loads the extracted data from a JSON file, updates the text and character count
    for each block in the PDF, and saves the updated data back to the JSON file.

    :param pdf_path: Path to the input PDF file.
    :param json_path: Path to the JSON file with extracted data.
    """
    # Load the JSON data
    with open(json_path, "r", encoding="utf-8") as json_file:
        extracted_data = json.load(json_file)

    # Create a document object
    doc = fitz.open(pdf_path)

    # Iterate through the extracted data and translate the text
    for page_num, page_data in extracted_data.items():

        # Get the page
        page = doc.load_page(int(page_num) - 1)  # or page = doc[i]
        blocks = page.get_text("blocks")
        for b,block in zip(blocks,page_data):
            text = b[4]
            block["text"] = text
            block["Character_count"] = len(text)


    # Save the extracted data to a JSON file
    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(extracted_data, json_file, ensure_ascii=False, indent=4)

    # Close the document
    doc.close()

    print(f"Text and character count added. Data saved to {json_path}.")

def translate_and_insert_newlinesk(pdf_path, json_path, dest_language='kn'):
    """
    This function loads extracted text from a JSON file, translates the text,
    inserts newlines at appropriate positions, and updates the JSON file with the translated text
    and character counts.

    :param pdf_path: Path to the input PDF file (unused here, kept for context).
    :param json_path: Path to the input/output JSON file.
    :param dest_language: The language code to translate to (default is Kannada 'kn').
    """
    # Initialize the translator
    translator = Translator()

    # Function to translate text
    def translate_text(text, dest_language='kn'):  # Change 'kn' to your desired language code
        try:
            translated = translator.translate(text, dest=dest_language)
            print("Translated:", translated.text)
            return (translated.text)
        except Exception as e:
            print(f"Error in translation: {e}")
            return text

    def insert_newlines(text, translated_text):
        # Find positions of '\n' in the original text
        newline_positions = [pos for pos, char in enumerate(text) if char == '\n']

        # Adjust positions for translated text
        adjusted_positions = []
        offset = 0
        for pos in newline_positions:
            while pos + offset < len(translated_text) and translated_text[pos + offset] != ' ':
                offset += 1
            adjusted_positions.append(pos + offset)

        # Insert '\n' at the adjusted positions
        for pos in reversed(adjusted_positions):  # reversed to avoid messing up positions
            translated_text = translated_text[:pos] + '\n' + translated_text[pos:]

        return translated_text


    # Load the JSON data
    with open(json_path, "r", encoding="utf-8") as json_file:
        extracted_data = json.load(json_file)

    # Iterate through the extracted data and translate the text
    for page_num, page_data in extracted_data.items():
        for block in page_data:
            original_text = block.get("text", "")
            c1 = block.get("Character_count", len(original_text)) 
            f = block.get("IniFontsize", 12)
            translated_text = translate_text(original_text, dest_language='kn')  # Change 'kn' to your desired language code
            block["translated_text"] = translated_text
            block["translated_character_count"] = len(translated_text)
            #block["Translated_text"] = insert_newlines(block["text"], translated_text)
            if c1 < len(translated_text):
                block["Font_Size"] = len(translated_text) * f / c1
            else:
                block["Font_Size"] = f

    # Save the translated data back to the same JSON file
    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(extracted_data, json_file, ensure_ascii=False, indent=4)
    
    print(f"Translation and updates complete. Data saved to {json_path}.")

def translate_and_insert_newlinesh(pdf_path, json_path, dest_language='hi'):
    """
    This function loads extracted text from a JSON file, translates the text,
    inserts newlines at appropriate positions, and updates the JSON file with the translated text
    and character counts.

    :param pdf_path: Path to the input PDF file (unused here, kept for context).
    :param json_path: Path to the input/output JSON file.
    :param dest_language: The language code to translate to (default is Kannada 'kn').
    """
    # Initialize the translator
    translator = Translator()

    # Function to translate text
    def translate_text(text, dest_language='hi'):  # Change 'kn' to your desired language code
        try:
            translated = translator.translate(text, dest=dest_language)
            print("Translated:", translated.text)
            return (translated.text)
        except Exception as e:
            print(f"Error in translation: {e}")
            return text

    def insert_newlines(text, translated_text):
        # Find positions of '\n' in the original text
        newline_positions = [pos for pos, char in enumerate(text) if char == '\n']

        # Adjust positions for translated text
        adjusted_positions = []
        offset = 0
        for pos in newline_positions:
            while pos + offset < len(translated_text) and translated_text[pos + offset] != ' ':
                offset += 1
            adjusted_positions.append(pos + offset)

        # Insert '\n' at the adjusted positions
        for pos in reversed(adjusted_positions):  # reversed to avoid messing up positions
            translated_text = translated_text[:pos] + '\n' + translated_text[pos:]

        return translated_text


    # Load the JSON data
    with open(json_path, "r", encoding="utf-8") as json_file:
        extracted_data = json.load(json_file)

    # Iterate through the extracted data and translate the text
    for page_num, page_data in extracted_data.items():
        for block in page_data:
            original_text = block.get("text", "")
            c1 = block.get("Character_count", len(original_text)) 
            f = block.get("IniFontsize", 12)
            translated_text = translate_text(original_text, dest_language='hi')  # Change 'kn' to your desired language code
            block["translated_text"] = translated_text
            block["translated_character_count"] = len(translated_text)
            #block["Translated_text"] = insert_newlines(block["text"], translated_text)
            if c1 < len(translated_text):
                block["Font_Size"] = len(translated_text) * f / c1
            else:
                block["Font_Size"] = f

    # Save the translated data back to the same JSON file
    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(extracted_data, json_file, ensure_ascii=False, indent=4)
    
    print(f"Translation and updates complete. Data saved to {json_path}.")


def translate_and_insert_newlinesr(pdf_path, json_path, dest_language='ru'):
    """
    This function loads extracted text from a JSON file, translates the text,
    inserts newlines at appropriate positions, and updates the JSON file with the translated text
    and character counts.

    :param pdf_path: Path to the input PDF file (unused here, kept for context).
    :param json_path: Path to the input/output JSON file.
    :param dest_language: The language code to translate to (default is Kannada 'kn').
    """
    # Initialize the translator
    translator = Translator()

    # Function to translate text
    def translate_text(text, dest_language='ru'):  # Change 'kn' to your desired language code
        try:
            translated = translator.translate(text, dest=dest_language)
            print("Translated:", translated.text)
            return (translated.text)
        except Exception as e:
            print(f"Error in translation: {e}")
            return text

    def insert_newlines(text, translated_text):
        # Find positions of '\n' in the original text
        newline_positions = [pos for pos, char in enumerate(text) if char == '\n']

        # Adjust positions for translated text
        adjusted_positions = []
        offset = 0
        for pos in newline_positions:
            while pos + offset < len(translated_text) and translated_text[pos + offset] != ' ':
                offset += 1
            adjusted_positions.append(pos + offset)

        # Insert '\n' at the adjusted positions
        for pos in reversed(adjusted_positions):  # reversed to avoid messing up positions
            translated_text = translated_text[:pos] + '\n' + translated_text[pos:]

        return translated_text


    # Load the JSON data
    with open(json_path, "r", encoding="utf-8") as json_file:
        extracted_data = json.load(json_file)

    # Iterate through the extracted data and translate the text
    for page_num, page_data in extracted_data.items():
        for block in page_data:
            original_text = block.get("text", "")
            c1 = block.get("Character_count", len(original_text)) 
            f = block.get("IniFontsize", 12)
            translated_text = translate_text(original_text, dest_language='ru')  # Change 'kn' to your desired language code
            block["translated_text"] = translated_text
            block["translated_character_count"] = len(translated_text)
            #block["Translated_text"] = insert_newlines(block["text"], translated_text)
            if c1 < len(translated_text):
                block["Font_Size"] = len(translated_text) * f / c1
            else:
                block["Font_Size"] = f

    # Save the translated data back to the same JSON file
    with open(json_path, "w", encoding="utf-8") as json_file:
        json.dump(extracted_data, json_file, ensure_ascii=False, indent=4)
    
    print(f"Translation and updates complete. Data saved to {json_path}.")


def create_translated_pdf(json_path, output_pdf_path, font_path):
    """
    This function reads translated text from a JSON file and adds it to a PDF document
    with the specified font, saving the result as an incremental update.

    :param json_path: Path to the JSON file containing extracted text with coordinates.
    :param output_pdf_path: Path to the existing PDF file where text will be added.
    :param font_path: Path to the TTF font file (e.g., Noto Sans Kannada).
    """
    # Check if the font file exists
    if not os.path.isfile(font_path):
        raise FileNotFoundError(f"The font file was not found: {font_path}")

    # Load the JSON data
    with open(json_path, "r", encoding="utf-8") as json_file:
        extracted_data = json.load(json_file)

    # Create a new PDF document
    new_doc = fitz.open(output_pdf_path)

    # Iterate through the pages in the JSON data
    for page_num, page_data in extracted_data.items():
        # Load a page
        new_page = new_doc.load_page(int(page_num)-1)

        # Iterate through the text blocks on the current page
        for block in page_data:
            translated_text = block["translated_text"]
            coordinates = block["coordinates"]
            #origin = block["origin"]
            x0, y0, x1, y1 = coordinates[0], coordinates[1], coordinates[2], coordinates[3]

            # Set the font size based on the height of the bounding box
            font_size = block["IniFontsize"]

            rect = fitz.Rect(x0, y0, x1, y1)

            print(translated_text)

            # Draw translated text on the new page with the font file using insert_textbox
            new_page.insert_htmlbox(
                coordinates,
                translated_text,
                #align=0,
                #fontsize=font_size,
                #fontname='NotoSansKannada',
                #fontfile=noto_sans_kannada_path,
                #color=(0, 0, 0),
                #oc=0
            )#

    # Save the new PDF
    new_doc.saveIncr()

    # Close the document
    new_doc.close()

    print(f"Translated PDF saved to {output_pdf_path}.")


# pdf_path = r"ScannedPDF_Final/TextOp.pdf"
# output_json_path = r"extracted_text_with_coordinates.json"
# extract_text_with_coordinates(pdf_path, output_json_path)
# add_text_and_character_count(pdf_path, output_json_path)
# translate_and_insert_newlines(pdf_path, output_json_path, dest_language='kn')
# font_path = r"NotoSansKannada-VariableFont_wdth,wght.ttf"
# create_translated_pdf(output_json_path, output_pdf_path, font_path)